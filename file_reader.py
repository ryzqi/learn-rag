"""
文件读取器

该模块提供了用于读取多种格式文档的客户端类。
支持PDF、Markdown、TXT、DOCX格式文件的文本内容提取。
"""

from pathlib import Path
from typing import Dict, Callable

try:
    import markdown
except ImportError:
    markdown = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None


class FileReader:
    """文件读取器客户端

    用于读取多种格式的文档文件，支持PDF、Markdown、TXT、DOCX格式的文本内容提取。
    根据文件扩展名自动选择相应的处理方法。
    """

    def __init__(self):
        """初始化文件读取器

        设置支持的文件格式映射表，将文件扩展名映射到相应的处理方法。
        """
        # 支持的文件格式映射表
        self.supported_formats: Dict[str, Callable[[str], str]] = {
            ".pdf": self._read_pdf,
            ".docx": self._read_docx,
            ".md": self._read_markdown,
            ".txt": self._read_txt,
        }

    def read_file(self, file_path: str) -> str:
        """读取文件内容

        根据文件扩展名自动选择相应的处理方法，提取文件中的文本内容。

        Args:
            file_path: 文件路径，支持相对路径和绝对路径

        Returns:
            提取的文本内容字符串

        Raises:
            ValueError: 当文件路径为空、文件格式不支持或文件读取失败时
            FileNotFoundError: 当文件不存在时
        """
        # 参数验证
        if not file_path or not file_path.strip():
            raise ValueError("文件路径不能为空")

        # 转换为Path对象便于处理
        path_obj = Path(file_path)

        # 检查文件是否存在
        if not path_obj.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 检查是否为文件（而非目录）
        if not path_obj.is_file():
            raise ValueError(f"路径不是文件: {file_path}")

        # 获取文件扩展名（转为小写）
        file_extension = path_obj.suffix.lower()

        # 检查文件格式是否支持
        if file_extension not in self.supported_formats:
            supported_formats = ", ".join(self.supported_formats.keys())
            raise ValueError(
                f"不支持的文件格式: {file_extension}。支持的格式: {supported_formats}"
            )

        # 调用相应的处理方法
        try:
            handler = self.supported_formats[file_extension]
            content = handler(str(path_obj.absolute()))

            # 基本的文本清理
            if content:
                content = content.strip()

            return content or ""

        except Exception as e:
            # 如果是我们已知的异常类型，直接抛出
            if isinstance(e, (ValueError, FileNotFoundError)):
                raise e
            # 其他异常包装为ValueError
            raise ValueError(f"文件读取失败: {e}")

    def _read_pdf(self, file_path: str) -> str:
        """读取PDF文件内容

        Args:
            file_path: PDF文件路径

        Returns:
            提取的文本内容

        Raises:
            ValueError: 当文件读取失败或pdfplumber库未安装时
        """
        # 检查pdfplumber库是否可用
        if pdfplumber is None:
            raise ValueError("pdfplumber库未安装，请运行: pip install pdfplumber")

        try:
            # 使用pdfplumber打开PDF文件
            with pdfplumber.open(file_path) as pdf:
                # 检查PDF是否有页面
                if not pdf.pages:
                    return ""

                # 提取所有页面的文本内容
                text_content = []

                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        # 尝试多种文本提取方法
                        page_text = None

                        # 方法1：使用默认的extract_text()
                        try:
                            page_text = page.extract_text()
                        except Exception:
                            pass

                        # 方法2：如果默认方法失败，尝试使用不同的参数
                        if not page_text or not page_text.strip():
                            try:
                                page_text = page.extract_text(
                                    x_tolerance=3,
                                    y_tolerance=3,
                                    layout=True,
                                    x_density=7.25,
                                    y_density=13,
                                )
                            except Exception:
                                pass

                        # 方法3：如果仍然失败，尝试提取字符
                        if not page_text or not page_text.strip():
                            try:
                                chars = page.chars
                                if chars:
                                    page_text = "".join(
                                        [char.get("text", "") for char in chars]
                                    )
                            except Exception:
                                pass

                        if page_text:
                            # 清理页面文本
                            page_text = page_text.strip()
                            if page_text:
                                text_content.append(page_text)

                    except Exception as e:
                        # 如果单个页面提取失败，记录但继续处理其他页面
                        # 使用警告而不是打印，以便更好的错误处理
                        import warnings

                        warnings.warn(f"第{page_num}页文本提取失败: {e}", UserWarning)
                        continue

                # 合并所有页面的文本
                if text_content:
                    # 使用双换行符分隔页面内容
                    full_text = "\n\n".join(text_content)

                    # 基本的文本清理
                    # 规范化空白字符
                    import re

                    full_text = re.sub(
                        r"\n\s*\n\s*\n", "\n\n", full_text
                    )  # 合并多个空行
                    full_text = re.sub(r"[ \t]+", " ", full_text)  # 规范化空格和制表符

                    return full_text.strip()
                else:
                    # 如果没有提取到任何文本内容
                    return ""

        except Exception as e:
            # 处理各种PDF相关的异常
            error_message = str(e).lower()

            if "password" in error_message or "encrypted" in error_message:
                raise ValueError(f"PDF文件受密码保护，无法读取: {file_path}")
            elif "damaged" in error_message or "corrupt" in error_message:
                raise ValueError(f"PDF文件已损坏，无法读取: {file_path}")
            elif "not a pdf" in error_message or "invalid pdf" in error_message:
                raise ValueError(f"文件不是有效的PDF格式: {file_path}")
            else:
                # 其他未知错误
                raise ValueError(f"读取PDF文件失败: {e}")

    def _read_docx(self, file_path: str) -> str:
        """读取DOCX文件内容

        Args:
            file_path: DOCX文件路径

        Returns:
            提取的文本内容

        Raises:
            ValueError: 当文件读取失败或python-docx库未安装时
        """
        # 检查python-docx库是否可用
        if Document is None:
            raise ValueError("python-docx库未安装，请运行: pip install python-docx")

        try:
            # 使用python-docx打开DOCX文件
            doc = Document(file_path)

            # 提取文档内容
            text_content = []

            # 1. 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())

            # 2. 提取表格文本
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))

                if table_text:
                    # 在表格前后添加空行以便区分
                    text_content.append("")
                    text_content.extend(table_text)
                    text_content.append("")

            # 3. 提取页眉页脚文本
            try:
                # 提取页眉
                for section in doc.sections:
                    if section.header:
                        for paragraph in section.header.paragraphs:
                            if paragraph.text.strip():
                                text_content.insert(
                                    0, f"[页眉] {paragraph.text.strip()}"
                                )

                    # 提取页脚
                    if section.footer:
                        for paragraph in section.footer.paragraphs:
                            if paragraph.text.strip():
                                text_content.append(f"[页脚] {paragraph.text.strip()}")
            except Exception:
                # 如果页眉页脚提取失败，忽略但继续处理
                pass

            # 合并所有文本内容
            if text_content:
                # 使用换行符连接所有内容
                full_text = "\n".join(text_content)

                # 基本的文本清理
                import re

                # 规范化空白字符
                full_text = re.sub(r"\n\s*\n\s*\n", "\n\n", full_text)  # 合并多个空行
                full_text = re.sub(r"[ \t]+", " ", full_text)  # 规范化空格和制表符
                full_text = re.sub(r"\r\n", "\n", full_text)  # 统一换行符

                return full_text.strip()
            else:
                # 如果没有提取到任何文本内容
                return ""

        except Exception as e:
            # 处理各种DOCX相关的异常
            error_message = str(e).lower()

            if "password" in error_message or "encrypted" in error_message:
                raise ValueError(f"DOCX文件受密码保护，无法读取: {file_path}")
            elif "corrupt" in error_message or "damaged" in error_message:
                raise ValueError(f"DOCX文件已损坏，无法读取: {file_path}")
            elif "not a zip file" in error_message or "bad zipfile" in error_message:
                raise ValueError(f"文件不是有效的DOCX格式: {file_path}")
            elif "no such file" in error_message or "cannot find" in error_message:
                raise ValueError(f"DOCX文件不存在或无法访问: {file_path}")
            else:
                # 其他未知错误
                raise ValueError(f"读取DOCX文件失败: {e}")

    def _read_markdown(self, file_path: str) -> str:
        """读取Markdown文件内容

        Args:
            file_path: Markdown文件路径

        Returns:
            提取的文本内容

        Raises:
            ValueError: 当文件读取失败或markdown库未安装时
        """
        # 检查markdown库是否可用
        if markdown is None:
            raise ValueError("markdown库未安装，请运行: pip install markdown")

        # 首先读取原始文件内容
        try:
            # 尝试多种编码格式读取文件
            encodings = ["utf-8", "utf-8-sig", "gbk", "gb2312", "latin-1"]

            raw_content = None
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        raw_content = file.read()
                        break
                except UnicodeDecodeError:
                    continue

            if raw_content is None:
                raise ValueError(f"无法使用支持的编码格式读取文件: {file_path}")

            # 将Markdown转换为HTML
            html_content = markdown.markdown(raw_content)

            # 从HTML中提取纯文本内容
            # 移除HTML标签
            import re

            text_content = re.sub(r"<[^>]+>", "", html_content)

            # 清理多余的空白字符
            text_content = re.sub(r"\n\s*\n", "\n\n", text_content)  # 规范化换行
            text_content = re.sub(r"[ \t]+", " ", text_content)  # 规范化空格

            return text_content.strip()

        except Exception as e:
            if isinstance(e, ValueError):
                raise e
            raise ValueError(f"读取Markdown文件失败: {e}")

    def _read_txt(self, file_path: str) -> str:
        """读取TXT文件内容

        Args:
            file_path: TXT文件路径

        Returns:
            提取的文本内容

        Raises:
            ValueError: 当文件读取失败时
        """
        # 尝试多种编码格式读取文件
        encodings = ["utf-8", "utf-8-sig", "gbk", "gb2312", "latin-1"]

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as file:
                    content = file.read()
                    return content
            except UnicodeDecodeError:
                continue
            except Exception as e:
                raise ValueError(f"读取TXT文件失败: {e}")

        # 如果所有编码都失败，抛出异常
        raise ValueError(f"无法使用支持的编码格式读取文件: {file_path}")


# 创建全局服务实例，便于直接导入使用
file_reader_service = FileReader()
